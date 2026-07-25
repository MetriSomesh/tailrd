#!/usr/bin/env python3
"""Generate a professionally formatted DOCX resume from tailored JSON."""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

from app.engine.validators import load_json_safe, validate_or_exit

FONT_NAME = "Segoe UI"


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_font(run, name=None, size=11, bold=False, color=None):
    name = name or FONT_NAME
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = hex_to_rgb(color)
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=10.5, bold=True, color="1A3A5C")
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "1A3A5C",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_font(run, size=9.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2 + indent_level * 0.2)
    return p


def render_skills(doc, skills):
    if isinstance(skills, dict):
        for category, items in skills.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"{category}: ")
            set_font(run, size=9.5, bold=True, color="333333")
            run = p.add_run(", ".join(items))
            set_font(run, size=9.5, color="555555")
    elif isinstance(skills, list):
        for i in range(0, len(skills), 6):
            chunk = skills[i : i + 6]
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(chunk))
            set_font(run, size=9.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def generate_resume(json_path, output_path):
    data = load_json_safe(json_path)
    validate_or_exit(data, json_path)

    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except OSError as e:
            print(f"ERROR: Cannot create output directory {output_dir}: {e}", file=sys.stderr)
            sys.exit(1)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    immutable = data["immutable"]
    editable = data["editable"]

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.space_before = Pt(0)
    run = header.add_run(immutable["name"])
    set_font(run, size=15, bold=True, color="1A3A5C")

    contact = immutable["contact"]
    contact_parts = []
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("email"):
        contact_parts.append(contact["email"])

    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(" | ".join(contact_parts))
        set_font(run, size=8.5, color="555555")

    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    run = p.add_run(editable["about"])
    set_font(run, size=9.5)
    p.paragraph_format.space_after = Pt(2)

    if editable.get("skills"):
        add_section_heading(doc, "Skills")
        render_skills(doc, editable["skills"])

    add_section_heading(doc, "Experience")
    for exp in editable["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(exp["title"])
        set_font(run, size=10, bold=True, color="1A3A5C")
        run = p.add_run(f"  |  {exp['company']}")
        set_font(run, size=9.5, color="555555")

        if exp.get("location") or exp.get("dates"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(1)
            parts = []
            if exp.get("location"):
                parts.append(exp["location"])
            if exp.get("dates"):
                parts.append(exp["dates"])
            run = p2.add_run(" | ".join(parts))
            set_font(run, size=8.5, color="777777")

        for bullet in exp["bullets"]:
            add_bullet(doc, bullet)

    if editable.get("projects"):
        add_section_heading(doc, "Projects")
        for proj in editable["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(proj["title"])
            set_font(run, size=10, bold=True, color="1A3A5C")

            if proj.get("technologies"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1)
                run = p2.add_run(", ".join(proj["technologies"]))
                set_font(run, size=8.5, color="555555")

            desc_lines = proj["description"].split("\n")
            for line in desc_lines:
                line = line.strip().lstrip("•-* ")
                if line:
                    add_bullet(doc, line)

    add_section_heading(doc, "Education")
    for edu in immutable["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(edu["degree"])
        set_font(run, size=9.5, bold=True, color="1A3A5C")

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        run = p2.add_run(edu["institution"])
        set_font(run, size=9, color="555555")
        if edu.get("dates"):
            run = p2.add_run(f"  |  {edu['dates']}")
            set_font(run, size=8.5, color="777777")

    try:
        doc.save(output_path)
    except PermissionError:
        print(
            f"ERROR: Permission denied writing to {output_path}. Is the file open?", file=sys.stderr
        )
        return False
    except Exception as e:
        print(f"ERROR: Failed to save DOCX: {e}", file=sys.stderr)
        return False

    print(f"Resume saved to {output_path}")
    return True


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_docx.py <input.json> [output.docx]")
        print("\nIf output.docx is omitted, filename is derived from candidate name in JSON.")
        sys.exit(1)

    json_path = sys.argv[1]

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # Auto-derive from candidate name
        data = load_json_safe(json_path)
        name = data.get("immutable", {}).get("name", "Resume")
        output_path = name.replace(" ", "_") + "_Resume.docx"

    success = generate_resume(json_path, output_path)
    sys.exit(0 if success else 1)
