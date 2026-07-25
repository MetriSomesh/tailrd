#!/usr/bin/env python3
"""Validate that a generated DOCX resume is parsable by ATS systems.

ATS systems (Workday, Greenhouse, Lever) extract text from DOCX files.
This script simulates that extraction and checks that critical resume content
survives the process. If content is lost, it means the DOCX formatting is
hiding information from ATS parsers.

Usage:
    python validate_ats_parsability.py <resume.docx> <tailored.json>
    python validate_ats_parsability.py <resume.docx> <tailored.json> --json

Exit codes:
    0 = Parsable (all critical content extractable)
    1 = Content loss detected (details printed)
    2 = File error
"""

import json
import os
import re
import sys

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)

from app.engine.validators import load_json_safe, validate_or_exit


def extract_all_text_from_docx(docx_path):
    """Extract all visible text from a DOCX file, simulating ATS parsing.

    Extracts from:
    - Body paragraphs
    - Tables
    - Does NOT extract from: headers/footers, text boxes, images, comments

    Returns:
        str: All extracted text joined with newlines.
    """
    if not os.path.exists(docx_path):
        print(f"ERROR: DOCX file not found: {docx_path}", file=sys.stderr)
        sys.exit(2)

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"ERROR: Cannot open DOCX file: {e}", file=sys.stderr)
        sys.exit(2)

    text_parts = []

    # Extract from body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return '\n'.join(text_parts)


def check_content_presence(extracted_text, content_item, label):
    """Check if a content item is present in extracted text.

    Handles fuzzy matching — normalizes whitespace and case for comparison.

    Returns:
        dict: {'found': bool, 'label': str, 'content': str (truncated)}
    """
    extracted_lower = extracted_text.lower()
    extracted_normalized = re.sub(r'\s+', ' ', extracted_lower)

    content_lower = content_item.lower().strip()
    content_normalized = re.sub(r'\s+', ' ', content_lower)

    # Try exact substring match first
    found = content_normalized in extracted_normalized

    # If not found, try matching first 50 chars (handles minor formatting differences)
    if not found and len(content_normalized) > 50:
        found = content_normalized[:50] in extracted_normalized

    # If still not found, try keyword match (>70% of words present)
    if not found:
        content_words = set(content_normalized.split())
        content_words = {w for w in content_words if len(w) > 3}
        if content_words:
            present_count = sum(1 for w in content_words if w in extracted_normalized)
            word_coverage = present_count / len(content_words)
            found = word_coverage >= 0.7

    return {
        'found': found,
        'label': label,
        'content': content_item[:80] + ('...' if len(content_item) > 80 else ''),
    }


def validate_parsability(docx_path, resume_json_path):
    """Run full ATS parsability validation.

    Checks:
    1. Candidate name is extractable
    2. Phone number is extractable
    3. Email is extractable
    4. All skills are extractable
    5. All experience bullets are extractable
    6. All project descriptions are extractable
    7. Section headings are present
    8. Total word count is reasonable (no major content loss)

    Returns:
        dict: Parsability report with score and details.
    """
    # Load resume data
    resume_data = load_json_safe(resume_json_path)
    validate_or_exit(resume_data, resume_json_path)

    # Extract text from DOCX
    extracted_text = extract_all_text_from_docx(docx_path)

    if not extracted_text.strip():
        return {
            'parsability_pct': 0.0,
            'status': 'FAIL',
            'error': 'No text could be extracted from DOCX',
            'checks': [],
            'warnings': ['The DOCX appears empty or uses non-extractable formatting (text boxes, images)'],
        }

    immutable = resume_data['immutable']
    editable = resume_data['editable']

    checks = []
    warnings = []

    # 1. Check candidate name
    checks.append(check_content_presence(
        extracted_text, immutable['name'], 'Candidate Name'))

    # 2. Check phone
    phone = immutable.get('contact', {}).get('phone', '')
    if phone:
        checks.append(check_content_presence(extracted_text, phone, 'Phone Number'))

    # 3. Check email
    email = immutable.get('contact', {}).get('email', '')
    if email:
        checks.append(check_content_presence(extracted_text, email, 'Email Address'))

    # 4. Check skills
    skills = editable.get('skills', [])
    if isinstance(skills, dict):
        # Categorized dict — check a sample of skills from each category
        for category, items in skills.items():
            if isinstance(items, list):
                for skill in items[:3]:  # Check first 3 per category
                    checks.append(check_content_presence(
                        extracted_text, skill, f'Skill ({category})'))
    elif isinstance(skills, list):
        for skill in skills[:10]:  # Check first 10
            checks.append(check_content_presence(
                extracted_text, skill, 'Skill'))

    # 5. Check experience bullets
    for i, exp in enumerate(editable.get('experience', [])):
        for j, bullet in enumerate(exp.get('bullets', [])):
            checks.append(check_content_presence(
                extracted_text, bullet, f"Experience[{i}] Bullet[{j}]"))

    # 6. Check project content
    for i, proj in enumerate(editable.get('projects', [])):
        checks.append(check_content_presence(
            extracted_text, proj.get('title', ''), f"Project[{i}] Title"))
        desc = proj.get('description', '')
        # Check each line of multi-line descriptions
        for line in desc.split('\n'):
            line = line.strip().lstrip('•-* ')
            if line and len(line) > 15:
                checks.append(check_content_presence(
                    extracted_text, line, f"Project[{i}] Description"))

    # 7. Check section headings
    expected_sections = ['PROFESSIONAL SUMMARY', 'SKILLS', 'EXPERIENCE', 'PROJECTS', 'EDUCATION']
    for section in expected_sections:
        checks.append(check_content_presence(
            extracted_text, section, f"Section Heading: {section}"))

    # 8. Check education
    for edu in immutable.get('education', []):
        checks.append(check_content_presence(
            extracted_text, edu.get('degree', ''), 'Education Degree'))

    # Calculate parsability score
    total_checks = len(checks)
    passed_checks = sum(1 for c in checks if c['found'])
    failed_checks = [c for c in checks if not c['found']]

    parsability_pct = (passed_checks / total_checks * 100) if total_checks > 0 else 0

    # Generate warnings
    if any(not c['found'] for c in checks if 'Phone' in c['label'] or 'Email' in c['label']):
        warnings.append("Contact info may be in a header/footer — some ATS systems skip those.")

    if any(not c['found'] for c in checks if 'Skill' in c['label']):
        warnings.append("Some skills may not be extractable — check if they are in text boxes or images.")

    if parsability_pct < 80:
        warnings.append("Significant content loss detected. Consider simplifying DOCX formatting.")

    # Word count sanity check
    extracted_word_count = len(extracted_text.split())
    if extracted_word_count < 100:
        warnings.append(f"Very low word count extracted ({extracted_word_count} words). "
                       "Possible heavy content loss.")

    status = 'PASS' if parsability_pct >= 90 else 'WARN' if parsability_pct >= 70 else 'FAIL'

    return {
        'parsability_pct': round(parsability_pct, 1),
        'status': status,
        'total_checks': total_checks,
        'passed_checks': passed_checks,
        'failed_checks_count': len(failed_checks),
        'failed_items': [{'label': c['label'], 'content': c['content']} for c in failed_checks],
        'warnings': warnings,
        'extracted_word_count': extracted_word_count,
    }


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python validate_ats_parsability.py <resume.docx> <tailored.json> [--json]")
        print("\nValidates that ATS systems can extract all critical content from the DOCX.")
        sys.exit(1)

    docx_path = sys.argv[1]
    json_path = sys.argv[2]
    output_json = '--json' in sys.argv

    result = validate_parsability(docx_path, json_path)

    if output_json:
        print(json.dumps(result, indent=2))
    else:
        print(f"\nATS Parsability: {result['parsability_pct']}% ({result['status']})")
        print(f"Checks: {result['passed_checks']}/{result['total_checks']} passed")
        print(f"Extracted words: {result['extracted_word_count']}")

        if result['failed_items']:
            print(f"\nMissing content ({result['failed_checks_count']} items):")
            for item in result['failed_items'][:10]:
                print(f"  - [{item['label']}] {item['content']}")
            if len(result['failed_items']) > 10:
                print(f"  ... and {len(result['failed_items']) - 10} more")

        if result['warnings']:
            print("\nWarnings:")
            for w in result['warnings']:
                print(f"  - {w}")

    # Exit code based on status
    if result['status'] == 'PASS':
        sys.exit(0)
    else:
        sys.exit(1)
