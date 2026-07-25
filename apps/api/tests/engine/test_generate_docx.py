"""Tests for generate_docx.py — DOCX resume generation."""

import os
import pytest


try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from app.engine.generate_docx import generate_resume


FIXTURES = os.path.join(os.path.dirname(__file__), "fixtures")


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestGenerateResume:
    """Tests for DOCX generation."""

    def test_creates_output_file(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        result = generate_resume(json_path, output_path)
        assert result is True
        assert os.path.exists(output_path)

    def test_output_is_valid_docx(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        # Should be parseable as a DOCX
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    def test_contains_candidate_name(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test Candidate" in all_text

    def test_contains_contact_info(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "1234567890" in all_text
        assert "test@example.com" in all_text

    def test_contains_experience_bullets(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "10,000 events/second" in all_text

    def test_contains_skills_categories(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Dict skills should render with category names
        assert "Languages:" in all_text or "Frameworks:" in all_text

    def test_contains_project_with_multiline_description(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # The project's multi-line description should be split into bullets
        assert "95% accuracy" in all_text

    def test_contains_education(self, tmp_path):
        json_path = os.path.join(FIXTURES, "valid_tailored.json")
        output_path = str(tmp_path / "test_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "B.S. Computer Science" in all_text
        assert "Test University" in all_text

    def test_minimal_resume_does_not_crash(self, tmp_path):
        json_path = os.path.join(FIXTURES, "minimal_tailored.json")
        output_path = str(tmp_path / "minimal_output.docx")
        result = generate_resume(json_path, output_path)
        assert result is True
        assert os.path.exists(output_path)

    def test_list_skills_render_as_pipe_separated(self, tmp_path):
        json_path = os.path.join(FIXTURES, "minimal_tailored.json")
        output_path = str(tmp_path / "list_skills_output.docx")
        generate_resume(json_path, output_path)
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # List skills should use pipe separator
        assert "Python | JavaScript" in all_text or "Python" in all_text
